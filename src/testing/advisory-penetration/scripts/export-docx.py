#!/usr/bin/env python3
"""
Export all advisory penetration test transcripts as Word (.docx) files.
Reads transcript.md from each scenario folder and saves a formatted .docx.
"""

import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Paths ─────────────────────────────────────────────────

SCRIPT_DIR = Path(__file__).parent
PENETRATION_DIR = SCRIPT_DIR.parent
TRANSCRIPTS_DIR = PENETRATION_DIR / "transcripts"
OUTPUT_DIR = PENETRATION_DIR / "transcripts-docx"

# ── Colour palette ────────────────────────────────────────

COLOR_TITLE      = RGBColor(0x1A, 0x1A, 0x2E)   # dark navy
COLOR_HEADING    = RGBColor(0x16, 0x21, 0x3E)   # navy
COLOR_SCORE      = RGBColor(0x0F, 0x3A, 0x60)   # dark blue
COLOR_USER       = RGBColor(0x1A, 0x53, 0x76)   # blue
COLOR_ASSISTANT  = RGBColor(0x1B, 0x5E, 0x20)   # dark green
COLOR_FLAG_HIGH  = RGBColor(0xE6, 0x5C, 0x00)   # orange
COLOR_FLAG_CRIT  = RGBColor(0xC6, 0x28, 0x28)   # red
COLOR_FLAG_OK    = RGBColor(0x2E, 0x7D, 0x32)   # green
COLOR_MUTED      = RGBColor(0x55, 0x55, 0x55)   # grey
COLOR_BLACK      = RGBColor(0x00, 0x00, 0x00)


# ── Document styling ──────────────────────────────────────

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, color=COLOR_HEADING):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level > 1 else 16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    size = {1: 16, 2: 13, 3: 11}.get(level, 10)
    set_font(run, size=size, bold=True, color=color)
    return p


def add_para(doc, text="", bold=False, italic=False, color=COLOR_BLACK, size=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_kv(doc, key, value, key_color=COLOR_MUTED, val_color=COLOR_BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    k = p.add_run(f"{key}: ")
    set_font(k, size=9, bold=True, color=key_color)
    v = p.add_run(value)
    set_font(v, size=9, color=val_color)
    return p


def add_divider(doc):
    p = doc.add_paragraph("─" * 72)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_font(run, size=8, color=COLOR_MUTED)


# ── Score bar ─────────────────────────────────────────────

def score_bar(score: int, width=20) -> str:
    filled = round(score / 100 * width)
    empty = width - filled
    return f"{'█' * filled}{'░' * empty} {score}/100"


def score_color(score: int) -> RGBColor:
    if score >= 80:
        return RGBColor(0x1B, 0x5E, 0x20)
    elif score >= 65:
        return RGBColor(0xE6, 0x5C, 0x00)
    else:
        return RGBColor(0xC6, 0x28, 0x28)


# ── Markdown parser ───────────────────────────────────────

def parse_transcript(md: str) -> dict:
    """Parse transcript.md into structured data."""
    result = {
        "title": "",
        "meta": {},
        "scenario_context": "",
        "scores": {},
        "red_flags": [],
        "strengths": [],
        "turns": [],
    }

    lines = md.splitlines()
    i = 0

    # Title (first # heading)
    while i < len(lines):
        line = lines[i]
        if line.startswith("# "):
            result["title"] = line[2:].strip()
            i += 1
            break
        i += 1

    # Meta key-value pairs (**Key:** Value)
    while i < len(lines):
        line = lines[i]
        m = re.match(r"\*\*(.+?):\*\*\s*(.*)", line)
        if m:
            result["meta"][m.group(1)] = m.group(2).strip()
        elif line.startswith("> "):
            result["scenario_context"] = line[2:].strip()
        elif line.startswith("## Evaluation"):
            break
        i += 1

    # Scores table
    while i < len(lines):
        line = lines[i]
        if line.startswith("## Red Flags"):
            break
        m = re.match(r"\|\s*(.+?)\s*\|\s*(\d+)/100\s*\|", line)
        if m and m.group(1) not in ("Dimension", "---"):
            result["scores"][m.group(1)] = int(m.group(2))
        i += 1

    # Red flags
    while i < len(lines):
        line = lines[i]
        if line.startswith("## Strengths"):
            break
        if line.startswith("- "):
            result["red_flags"].append(line[2:].strip())
        i += 1

    # Strengths
    while i < len(lines):
        line = lines[i]
        if line.startswith("## Conversation"):
            break
        if line.startswith("- "):
            result["strengths"].append(line[2:].strip())
        i += 1

    # Conversation turns
    current_turn = None
    current_role = None
    current_text_lines = []

    def flush_turn_text():
        if current_turn is not None and current_role and current_text_lines:
            text = "\n".join(current_text_lines).strip()
            current_turn[current_role] = text

    while i < len(lines):
        line = lines[i]

        if re.match(r"^### Turn \d+", line):
            flush_turn_text()
            current_turn = {"num": line.strip(), "user": "", "assistant": "", "meta": "", "flags": []}
            result["turns"].append(current_turn)
            current_role = None
            current_text_lines = []

        elif line.strip() == "**User:**" and current_turn is not None:
            flush_turn_text()
            current_role = "user"
            current_text_lines = []

        elif re.match(r"^\*\*Assistant\*\*", line) and current_turn is not None:
            flush_turn_text()
            current_turn["meta"] = line.strip()
            current_role = "assistant"
            current_text_lines = []

        elif line.startswith("**Turn Scores:**") and current_turn is not None:
            flush_turn_text()
            current_turn["scores_line"] = line.strip()
            current_role = None
            current_text_lines = []

        elif line.startswith("**Flags:**") and current_turn is not None:
            current_role = "flags"
            current_text_lines = []

        elif current_role == "flags" and line.startswith("- "):
            current_turn["flags"].append(line[2:].strip())

        elif re.match(r"^---+$", line) and current_turn is not None:
            flush_turn_text()
            current_role = None
            current_text_lines = []

        elif current_role in ("user", "assistant"):
            # Strip blockquote markers
            stripped = re.sub(r"^>\s?", "", line)
            current_text_lines.append(stripped)

        i += 1

    flush_turn_text()
    return result


# ── Build Word document ───────────────────────────────────

def build_docx(data: dict, scenario_id: str) -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Title ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"[{scenario_id}] {data['title']}")
    set_font(run, size=18, bold=True, color=COLOR_TITLE)

    # ── Meta ───────────────────────────────────────────────
    for key, val in data["meta"].items():
        add_kv(doc, key, val)

    if data["scenario_context"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"Context: {data['scenario_context']}")
        set_font(run, size=9, italic=True, color=COLOR_MUTED)

    add_divider(doc)

    # ── Scores ─────────────────────────────────────────────
    add_heading(doc, "Evaluation Scores", level=2, color=COLOR_SCORE)
    for dim, score in data["scores"].items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label = p.add_run(f"{dim:<38}")
        set_font(label, size=9, color=COLOR_MUTED)
        bar = p.add_run(score_bar(score))
        set_font(bar, size=9, bold=True, color=score_color(score))

    add_divider(doc)

    # ── Red Flags ──────────────────────────────────────────
    if data["red_flags"]:
        add_heading(doc, f"Red Flags ({len(data['red_flags'])})", level=2, color=COLOR_FLAG_HIGH)
        for flag in data["red_flags"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            # Determine severity from emoji
            if "🔴" in flag or "CRITICAL" in flag.upper():
                color = COLOR_FLAG_CRIT
            elif "🟠" in flag or "[HIGH]" in flag.upper():
                color = COLOR_FLAG_HIGH
            else:
                color = COLOR_MUTED
            clean = re.sub(r"[🔴🟠🟡✅⚠️]", "", flag).strip()
            run = p.add_run(clean)
            set_font(run, size=9, color=color)

    # ── Strengths ──────────────────────────────────────────
    if data["strengths"]:
        add_heading(doc, "Strengths", level=2, color=COLOR_FLAG_OK)
        for s in data["strengths"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            clean = re.sub(r"[✅⚠️🔴🟠🟡]", "", s).strip()
            run = p.add_run(clean)
            set_font(run, size=9, color=COLOR_FLAG_OK)

    add_divider(doc)

    # ── Conversation ───────────────────────────────────────
    add_heading(doc, "Conversation", level=2)

    for turn in data["turns"]:
        # Turn header
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(turn["num"])
        set_font(run, size=10, bold=True, color=COLOR_HEADING)

        # User message
        if turn.get("user"):
            add_para(doc, "USER", bold=True, color=COLOR_USER, size=9)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(turn["user"].strip())
            set_font(run, size=9.5, color=COLOR_BLACK)

        # Assistant meta line
        if turn.get("meta"):
            clean_meta = re.sub(r"\*\*|\*", "", turn["meta"])
            add_para(doc, clean_meta, italic=True, color=COLOR_MUTED, size=8)

        # Assistant message
        if turn.get("assistant"):
            add_para(doc, "ASSISTANT", bold=True, color=COLOR_ASSISTANT, size=9)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(turn["assistant"].strip())
            set_font(run, size=9.5, color=COLOR_BLACK)

        # Turn scores
        if turn.get("scores_line"):
            clean = re.sub(r"\*\*Turn Scores:\*\*\s*", "", turn["scores_line"])
            add_para(doc, f"Scores: {clean}", color=COLOR_MUTED, size=8)

        # Turn flags
        for flag in turn.get("flags", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            clean = re.sub(r"[🔴🟠🟡✅⚠️]", "", flag).strip()
            run = p.add_run(f"⚑ {clean}")
            color = COLOR_FLAG_CRIT if "CRITICAL" in flag.upper() else COLOR_FLAG_HIGH
            set_font(run, size=8, color=color)

        add_divider(doc)

    return doc


# ── Main ──────────────────────────────────────────────────

def main():
    # Find the most recent date folder
    if not TRANSCRIPTS_DIR.exists():
        print(f"ERROR: Transcripts directory not found: {TRANSCRIPTS_DIR}")
        sys.exit(1)

    date_dirs = sorted(TRANSCRIPTS_DIR.iterdir(), reverse=True)
    date_dirs = [d for d in date_dirs if d.is_dir()]
    if not date_dirs:
        print("ERROR: No transcript date folders found.")
        sys.exit(1)

    date_dir = date_dirs[0]
    print(f"Reading transcripts from: {date_dir}")

    scenario_dirs = sorted(date_dir.iterdir())
    scenario_dirs = [d for d in scenario_dirs if d.is_dir()]
    print(f"Found {len(scenario_dirs)} scenarios.")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Saving .docx files to: {OUTPUT_DIR}\n")

    ok = 0
    errors = 0

    for scenario_dir in scenario_dirs:
        scenario_id = scenario_dir.name
        md_file = scenario_dir / "transcript.md"

        if not md_file.exists():
            print(f"  [{scenario_id}] SKIP — no transcript.md")
            errors += 1
            continue

        try:
            md = md_file.read_text(encoding="utf-8")
            data = parse_transcript(md)
            doc = build_docx(data, scenario_id)

            out_file = OUTPUT_DIR / f"{scenario_id}.docx"
            doc.save(out_file)
            print(f"  [{scenario_id}] ✓ {data['title'][:60]}")
            ok += 1
        except Exception as e:
            print(f"  [{scenario_id}] ERROR: {e}")
            errors += 1

    print(f"\n{'='*60}")
    print(f"Done. {ok} files saved, {errors} errors.")
    print(f"Output: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
