#!/usr/bin/env python3
"""
Merge all PL transcripts into one .docx and all EN/mixed transcripts into another.
"""

import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).parent
PENETRATION_DIR = SCRIPT_DIR.parent
TRANSCRIPTS_DIR = PENETRATION_DIR / "transcripts"
OUTPUT_DIR = PENETRATION_DIR / "transcripts-docx"

COLOR_TITLE     = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_HEADING   = RGBColor(0x16, 0x21, 0x3E)
COLOR_SCORE     = RGBColor(0x0F, 0x3A, 0x60)
COLOR_USER      = RGBColor(0x1A, 0x53, 0x76)
COLOR_ASSISTANT = RGBColor(0x1B, 0x5E, 0x20)
COLOR_FLAG_HIGH = RGBColor(0xE6, 0x5C, 0x00)
COLOR_FLAG_CRIT = RGBColor(0xC6, 0x28, 0x28)
COLOR_FLAG_OK   = RGBColor(0x2E, 0x7D, 0x32)
COLOR_MUTED     = RGBColor(0x55, 0x55, 0x55)
COLOR_BLACK     = RGBColor(0x00, 0x00, 0x00)
COLOR_SEPARATOR = RGBColor(0xCC, 0xCC, 0xCC)


# ── Helpers ───────────────────────────────────────────────

def set_font(run, size=10, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    return br


def insert_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_para(doc, text="", bold=False, italic=False, color=COLOR_BLACK, size=10, space_before=0, space_after=4, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_divider(doc, char="─", width=70, color=COLOR_SEPARATOR):
    p = doc.add_paragraph(char * width)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_font(run, size=8, color=color)


def score_bar(score, width=18):
    filled = round(score / 100 * width)
    return f"{'█' * filled}{'░' * (width - filled)} {score}/100"


def score_color(score):
    if score >= 80:
        return COLOR_FLAG_OK
    elif score >= 65:
        return COLOR_FLAG_HIGH
    else:
        return COLOR_FLAG_CRIT


# ── Parser ────────────────────────────────────────────────

def parse_transcript(md):
    result = {"title": "", "meta": {}, "context": "", "scores": {}, "red_flags": [], "strengths": [], "turns": []}
    lines = md.splitlines()
    i = 0

    while i < len(lines):
        if lines[i].startswith("# "):
            result["title"] = lines[i][2:].strip()
            i += 1
            break
        i += 1

    while i < len(lines):
        line = lines[i]
        m = re.match(r"\*\*(.+?):\*\*\s*(.*)", line)
        if m:
            result["meta"][m.group(1)] = m.group(2).strip()
        elif line.startswith("> "):
            result["context"] = line[2:].strip()
        elif line.startswith("## Evaluation"):
            break
        i += 1

    while i < len(lines):
        line = lines[i]
        if line.startswith("## Red Flags"):
            break
        m = re.match(r"\|\s*(.+?)\s*\|\s*(\d+)/100\s*\|", line)
        if m and m.group(1) not in ("Dimension", "---"):
            result["scores"][m.group(1)] = int(m.group(2))
        i += 1

    while i < len(lines):
        line = lines[i]
        if line.startswith("## Strengths"):
            break
        if line.startswith("- "):
            result["red_flags"].append(line[2:].strip())
        i += 1

    while i < len(lines):
        line = lines[i]
        if line.startswith("## Conversation"):
            break
        if line.startswith("- "):
            result["strengths"].append(line[2:].strip())
        i += 1

    current_turn = None
    current_role = None
    buf = []

    def flush():
        if current_turn and current_role and buf:
            current_turn[current_role] = "\n".join(buf).strip()

    while i < len(lines):
        line = lines[i]
        if re.match(r"^### Turn \d+", line):
            flush()
            current_turn = {"num": line.strip(), "user": "", "assistant": "", "meta": "", "flags": [], "scores_line": ""}
            result["turns"].append(current_turn)
            current_role = None
            buf = []
        elif line.strip() == "**User:**" and current_turn is not None:
            flush(); current_role = "user"; buf = []
        elif re.match(r"^\*\*Assistant\*\*", line) and current_turn is not None:
            flush(); current_turn["meta"] = line.strip(); current_role = "assistant"; buf = []
        elif line.startswith("**Turn Scores:**") and current_turn is not None:
            flush(); current_turn["scores_line"] = line.strip(); current_role = None; buf = []
        elif line.startswith("**Flags:**") and current_turn is not None:
            current_role = "flags"; buf = []
        elif current_role == "flags" and line.startswith("- "):
            current_turn["flags"].append(line[2:].strip())
        elif re.match(r"^---+$", line) and current_turn is not None:
            flush(); current_role = None; buf = []
        elif current_role in ("user", "assistant"):
            buf.append(re.sub(r"^>\s?", "", line))
        i += 1
    flush()
    return result


# ── Write one scenario into doc ───────────────────────────

def write_scenario(doc, data, scenario_id, index, total):
    # Scenario header bar
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[{index}/{total}]  {scenario_id}")
    set_font(run, size=11, bold=True, color=COLOR_MUTED)

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(data["title"])
    set_font(run, size=15, bold=True, color=COLOR_TITLE)

    # Meta
    meta_items = []
    for k in ("Category", "Locale", "Overall Score", "Duration"):
        if k in data["meta"]:
            meta_items.append(f"{k}: {data['meta'][k]}")
    if meta_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("  ·  ".join(meta_items))
        set_font(run, size=8.5, color=COLOR_MUTED)

    # Context
    if data["context"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"Scenario: {data['context']}")
        set_font(run, size=9, italic=True, color=COLOR_MUTED)

    add_divider(doc)

    # Scores
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("SCORES")
    set_font(run, size=8, bold=True, color=COLOR_SCORE)

    for dim, score in data["scores"].items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        label = p.add_run(f"{dim:<38}")
        set_font(label, size=8.5, color=COLOR_MUTED)
        bar = p.add_run(score_bar(score))
        set_font(bar, size=8.5, bold=True, color=score_color(score))

    # Red flags
    if data["red_flags"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"RED FLAGS ({len(data['red_flags'])})")
        set_font(run, size=8, bold=True, color=COLOR_FLAG_HIGH)
        for flag in data["red_flags"]:
            p2 = doc.add_paragraph(style="List Bullet")
            p2.paragraph_format.space_after = Pt(1)
            clean = re.sub(r"[🔴🟠🟡✅⚠️]", "", flag).strip()
            color = COLOR_FLAG_CRIT if ("CRITICAL" in flag.upper() or "🔴" in flag) else COLOR_FLAG_HIGH
            run2 = p2.add_run(clean)
            set_font(run2, size=8.5, color=color)

    # Strengths
    if data["strengths"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("STRENGTHS")
        set_font(run, size=8, bold=True, color=COLOR_FLAG_OK)
        for s in data["strengths"]:
            p2 = doc.add_paragraph(style="List Bullet")
            p2.paragraph_format.space_after = Pt(1)
            clean = re.sub(r"[✅⚠️🔴🟠🟡]", "", s).strip()
            run2 = p2.add_run(clean)
            set_font(run2, size=8.5, color=COLOR_FLAG_OK)

    add_divider(doc)

    # Conversation
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("CONVERSATION")
    set_font(run, size=8, bold=True, color=COLOR_HEADING)

    for turn in data["turns"]:
        # Turn number
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(turn["num"])
        set_font(run, size=9.5, bold=True, color=COLOR_HEADING)

        # User
        if turn.get("user"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run("USER")
            set_font(run, size=8, bold=True, color=COLOR_USER)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(5)
            p2.paragraph_format.left_indent = Inches(0.15)
            run2 = p2.add_run(turn["user"].strip())
            set_font(run2, size=9.5, color=COLOR_BLACK)

        # Assistant meta
        if turn.get("meta"):
            clean = re.sub(r"\*\*|\*", "", turn["meta"])
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(clean)
            set_font(run, size=7.5, italic=True, color=COLOR_MUTED)

        # Assistant
        if turn.get("assistant"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run("ASSISTANT")
            set_font(run, size=8, bold=True, color=COLOR_ASSISTANT)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(5)
            p2.paragraph_format.left_indent = Inches(0.15)
            run2 = p2.add_run(turn["assistant"].strip())
            set_font(run2, size=9.5, color=COLOR_BLACK)

        # Turn scores + flags
        if turn.get("scores_line"):
            clean = re.sub(r"\*\*Turn Scores:\*\*\s*", "", turn["scores_line"])
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"Scores: {clean}")
            set_font(run, size=7.5, color=COLOR_MUTED)

        for flag in turn.get("flags", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            clean = re.sub(r"[🔴🟠🟡✅⚠️]", "", flag).strip()
            color = COLOR_FLAG_CRIT if "CRITICAL" in flag.upper() else COLOR_FLAG_HIGH
            run = p.add_run(f"⚑ {clean}")
            set_font(run, size=8, color=color)


# ── Cover page ────────────────────────────────────────────

def add_cover(doc, title, subtitle, count, date_str):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = 1  # center
    run = p.add_run(title)
    set_font(run, size=28, bold=True, color=COLOR_TITLE)

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(subtitle)
    set_font(run, size=14, color=COLOR_MUTED)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(f"{count} scenarios  ·  Generated: {date_str}")
    set_font(run, size=10, italic=True, color=COLOR_MUTED)

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run("ETAP 7.5 Advisory Penetration Test System")
    set_font(run, size=9, color=COLOR_SEPARATOR)

    insert_page_break(doc)


# ── Table of contents ─────────────────────────────────────

def add_toc(doc, entries):
    p = doc.add_paragraph()
    run = p.add_run("TABLE OF CONTENTS")
    set_font(run, size=13, bold=True, color=COLOR_TITLE)
    doc.add_paragraph()

    for idx, (sid, title, score) in enumerate(entries, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        num = p.add_run(f"{idx:>3}.  [{sid}]  ")
        set_font(num, size=8.5, color=COLOR_MUTED)
        t = p.add_run(title[:65])
        set_font(t, size=8.5, color=COLOR_BLACK)
        sc = p.add_run(f"  {score}/100")
        set_font(sc, size=8.5, bold=True, color=score_color(score))

    insert_page_break(doc)


# ── Main ──────────────────────────────────────────────────

def main():
    date_dirs = sorted(TRANSCRIPTS_DIR.iterdir(), reverse=True)
    date_dirs = [d for d in date_dirs if d.is_dir()]
    if not date_dirs:
        print("No transcript date folders found.")
        sys.exit(1)

    date_dir = date_dirs[0]
    date_str = date_dir.name
    print(f"Reading from: {date_dir}")

    scenario_dirs = sorted([d for d in date_dir.iterdir() if d.is_dir()])

    pl_entries = []
    en_entries = []

    for sd in scenario_dirs:
        md_file = sd / "transcript.md"
        if not md_file.exists():
            continue
        md = md_file.read_text(encoding="utf-8")
        data = parse_transcript(md)
        locale = data["meta"].get("Locale", "pl").strip().lower()
        score = int(data["meta"].get("Overall Score", "0").replace("/100", ""))
        entry = (sd.name, data["title"], score, data)

        if locale == "pl":
            pl_entries.append(entry)
        else:  # en, mixed → EN file
            en_entries.append(entry)

    print(f"PL: {len(pl_entries)} scenarios | EN/mixed: {len(en_entries)} scenarios")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    for lang, entries, filename in [
        ("PL", pl_entries, f"TRANSCRIPTS-PL-{date_str}.docx"),
        ("EN", en_entries, f"TRANSCRIPTS-EN-{date_str}.docx"),
    ]:
        print(f"\nBuilding {filename}...")
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.9)
            section.bottom_margin = Inches(0.9)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)

        lang_label = "Polish" if lang == "PL" else "English / Mixed"
        add_cover(
            doc,
            "Advisory Penetration Test",
            f"Transcripts — {lang_label}",
            len(entries),
            date_str,
        )

        toc_entries = [(sid, title, score) for sid, title, score, _ in entries]
        add_toc(doc, toc_entries)

        for idx, (sid, title, score, data) in enumerate(entries, 1):
            write_scenario(doc, data, sid, idx, len(entries))
            if idx < len(entries):
                insert_page_break(doc)
            print(f"  [{sid}] ✓ {title[:55]}")

        out = OUTPUT_DIR / filename
        doc.save(out)
        print(f"  Saved → {out}")

    print("\nDone.")


if __name__ == "__main__":
    main()
